"""Tests for the Phase 3 source parsers (txt/md, docx, html) and dispatch."""
from __future__ import annotations

from opennote.ingest.chunking import ChunkSpec
from opennote.ingest.parsers.docx import DocxParser
from opennote.ingest.parsers.html import HtmlParser, parse_url
from opennote.ingest.parsers.text import TextParser
from opennote.ingest.pipeline import find_source_files, get_parser_for_file


def test_text_parser_line_citations(tmp_path):
    p = tmp_path / "a.txt"
    p.write_text(
        "line one\nline two\nline three\nline four\n", encoding="utf-8"
    )
    chunks = TextParser().parse(p, ChunkSpec())
    assert chunks
    for c in chunks:
        assert 1 <= c.metadata["line_start"] <= c.metadata["line_end"]
        assert c.metadata["filename"] == "a.txt"


def test_text_parser_long_file_multiple_chunks(tmp_path):
    p = tmp_path / "long.txt"
    p.write_text("\n".join(f"sentence number {i} with padding content." for i in range(200)), encoding="utf-8")
    chunks = TextParser().parse(p, ChunkSpec(size=200, overlap=20))
    assert len(chunks) > 1


def test_docx_parser_heading_locators(tmp_path):
    from docx import Document

    p = tmp_path / "a.docx"
    d = Document()
    d.add_heading("Intro", level=1)
    d.add_paragraph("Body text under intro with several words.")
    d.add_heading("Second Section", level=2)
    d.add_paragraph("More body text here.")
    d.save(p)

    chunks = DocxParser().parse(p, ChunkSpec())
    headings = {c.metadata.get("heading") for c in chunks if "heading" in c.metadata}
    assert "Intro" in headings
    assert "Second Section" in headings


def test_html_parser_section_heading(tmp_path):
    p = tmp_path / "page.html"
    p.write_text(
        "<html><head><title>T</title></head><body>"
        "<h1>First Section</h1><p>Some content here.</p>"
        "<h2>Second Section</h2><p>More content here.</p>"
        "</body></html>",
        encoding="utf-8",
    )
    chunks = HtmlParser().parse(p, ChunkSpec())
    assert chunks
    assert any("heading" in c.metadata for c in chunks)


def test_find_source_files_multiformat(tmp_path):
    (tmp_path / "a.txt").write_text("x", encoding="utf-8")
    (tmp_path / "b.md").write_text("y", encoding="utf-8")
    (tmp_path / "c.docx").write_bytes(b"PK")
    (tmp_path / "d.html").write_text("z", encoding="utf-8")
    files = find_source_files(tmp_path)
    exts = {f.suffix for f in files}
    assert {".txt", ".md", ".docx", ".html"} <= exts


def test_parser_dispatch(tmp_path):
    assert isinstance(get_parser_for_file(tmp_path / "a.txt"), TextParser)
    assert isinstance(get_parser_for_file(tmp_path / "a.md"), TextParser)
    assert isinstance(get_parser_for_file(tmp_path / "a.markdown"), TextParser)
    assert isinstance(get_parser_for_file(tmp_path / "a.docx"), DocxParser)
    assert isinstance(get_parser_for_file(tmp_path / "a.html"), HtmlParser)
    assert isinstance(get_parser_for_file(tmp_path / "a.htm"), HtmlParser)
    assert isinstance(
        get_parser_for_file(tmp_path / "a.pdf", pdf_strategy="fallback"),
        __import__("opennote.ingest.parsers.pdf_fallback", fromlist=["FallbackPDFParser"]).FallbackPDFParser,
    )
    assert get_parser_for_file(tmp_path / "a.xyz") is None


def test_same_named_files_distinct_chunk_ids(tmp_path):
    # Two files named 'a.txt' in different directories must never collide (L12).
    d1 = tmp_path / "d1"
    d2 = tmp_path / "d2"
    d1.mkdir()
    d2.mkdir()
    a = d1 / "a.txt"
    b = d2 / "a.txt"
    a.write_text("content alpha in file one.", encoding="utf-8")
    b.write_text("content beta in file two.", encoding="utf-8")
    ca = TextParser().parse(a, ChunkSpec())
    cb = TextParser().parse(b, ChunkSpec())
    assert ca and cb
    assert {c.chunk_id for c in ca}.isdisjoint({c.chunk_id for c in cb})


def test_parse_url_filename_includes_path(monkeypatch):
    # URL chunks must be identified by host+path, not bare hostname (L30).
    html = "<html><body><h1>Title</h1><p>Body text here.</p></body></html>"
    import trafilatura

    monkeypatch.setattr(trafilatura, "fetch_url", lambda url: html)
    chunks = parse_url("https://example.com/docs/guide", ChunkSpec())
    assert chunks
    assert all(c.metadata["filename"] == "example.com/docs/guide" for c in chunks)